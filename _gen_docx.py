# -*- coding: utf-8 -*-
"""生成《大一新生班主任班级管理系统 使用说明.docx》，嵌入三端入口二维码。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
QR = os.path.join(BASE, "qr")
OUT = os.path.join(BASE, "使用说明_班主任班级管理系统.docx")
FONT = "微软雅黑"

doc = Document()

# 默认正体中文字体
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def set_cn(run):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: set_cn(r)
    return p
def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: set_cn(r)
    return p
def para(t, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; set_cn(r)
    return p
def bullet(t):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(t); set_cn(r)
    return p
def num(t):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(t); set_cn(r)
    return p

# ---------- 标题 ----------
title = doc.add_heading("大一新生班主任班级管理系统", level=0)
for r in title.runs: set_cn(r)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("使用说明（教师端 / 学生端 / 家长端）"); set_cn(r); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x55,0x55,0x55)

# ---------- 1 系统简介 ----------
h1("一、系统简介")
para("本系统参考机械制图作业管理系统架构，采用 Cloudflare Pages + D1 数据库 + Pages Functions 后端，免费部署、自带公网地址，教师电脑、学生手机、家长手机天然跨设备互通。系统包含三个终端：")
bullet("教师端：录入花名册（仅花名册内学生可使用系统），统揽全部学生，管理课表、上传期末成绩、查看每月学习情况记录与学生档案。")
bullet("学生端：用「姓名 + 学号」登录；填写个人档案（家庭/心理）、每月学习记录（学习状态/疑惑/心情/恋爱/生活，仅自己可见，教师与家长可看）；查看本人课表与成绩。")
bullet("家长端：用孩子的「姓名 + 学号」登录，只读查看自己孩子的学习、家庭、心理情况。")

# ---------- 2 部署 ----------
h1("二、部署步骤（一次性）")
para("1. 注册 Cloudflare 账号，新建 Pages 项目，连接本目录仓库。")
para("2. 构建输出目录为 class-mgr/（即本文件夹），框架选择“无/其他”。")
para("3. 在 Pages 项目设置中创建 D1 数据库（如 class-mgr-db），并在 wrangler.toml 的 [[d1_databases]] 中填写正确的 database_id 与 binding=DB。")
para("4. 设置环境变量/密钥 ADMIN（教师端密码，默认 teacher123，可在 wrangler.toml 修改）。")
para("5. 部署后执行一次建表：将 schema.sql 在 D1 控制台执行。")
para("6. 部署成功后得到公网域名（如 https://xxx.pages.dev）。如需把二维码指向该域名，修改 _gen_qr.py 顶部的 BASE_URL 后重跑即可。")

# ---------- 3 二维码 ----------
h1("三、二维码说明")
para("系统为三个终端各提供一个统一入口二维码（下图）：教师端、学生端、家长端分别扫码/访问进入对应入口页。学生端与家长端不再为每人单独生成二维码，统一使用对应入口码即可。")
para("二维码均在教师端「系统二维码」页实时生成（基于当前公网域名），可逐项下载或整页打印。")
doc.add_paragraph()

def add_qr(name, label):
    path = os.path.join(QR, name)
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label); set_cn(run); run.bold = True
        doc.add_picture(path, width=Inches(2.0))
        cap = doc.paragraphs[-1]; cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

h2("终端入口二维码（学生端、家长端各一个，统一管理）")
add_qr("entry_teacher.png", "教师端入口")
add_qr("entry_student.png", "学生端入口")
add_qr("entry_parent.png", "家长端入口")

# ---------- 4 教师端 ----------
h1("四、教师端使用")
num("用教师密码登录（默认 teacher123）。")
num("进入「花名册」：逐条添加或批量导入（格式 学号,姓名,班级，每行一条；或粘贴 JSON 数组）。仅花名册内的学生可登录系统。")
num("进入「系统二维码」：页面展示教师端、学生端、家长端三个统一入口码，下载或打印后分发；学生端与家长端各用一个入口码，登录时用「姓名 + 学号」。")
num("进入「课表管理」：选择学生，按“节次,课程,教室”逐行录入课表并保存。")
num("进入「成绩上传」：选择学生，按“课程,成绩,学分”逐行录入期末成绩并保存。")
num("进入「月记录总览」：查看全体学生每月学习情况记录汇总表。")
num("进入「档案查看」：选择学生查看其家庭与心理档案。")

# ---------- 5 学生端 ----------
h1("五、学生端使用")
num("用「姓名 + 学号」登录（与花名册中录入的一致）。")
num("「个人档案」：填写家庭情况（住址/监护人/经济/说明）与心理情况（压力/情绪/是否需关注/说明），保存后仅自己可见，教师与家长可查看。")
num("「每月学习记录」：填写当月学习状态、疑惑、心情、恋爱状态、生活；该记录仅自己可见，教师端与家长端可查看。可跨月多次记录。")
num("「我的课表 / 我的成绩」：查看教师已录入的本人课表与成绩。")

# ---------- 6 家长端 ----------
h1("六、家长端使用")
num("用孩子的「姓名 + 学号」登录（与花名册中录入的一致）。")
num("页面直接展示孩子的家庭情况、心理情况、课表、成绩与每月学习记录，只读不可改。")
num("家长端与学生端共用同一个家长入口二维码，扫码进入登录页即可。")

# ---------- 7 权限 ----------
h1("七、隐私与权限说明")
bullet("学生月记录：仅本人可见，教师端、家长端可查看（学生端“仅自己可见”指其他学生看不到）。")
bullet("档案：本人可写，教师与家长可读。")
bullet("课表/成绩：学生本人、家长、教师可读；成绩仅教师可上传。")
bullet("花名册外学生无法登录系统。")
bullet("所有数据存于 Cloudflare D1，按角色做服务端鉴权隔离，非前端明文。")

# ---------- 8 FAQ ----------
h1("八、常见问题")
bullet("学生/家长登录失败：请确认输入的「姓名」与「学号」与教师端花名册完全一致（含空格）。系统不再使用密码，凭姓名 + 学号即可登录。")
bullet("家长入口二维码失效：统一入口码长期有效，若域名变更请用教师端「系统二维码」页重新下载打印。")
bullet("跨设备不同步：确认教师端、学生端、家长端访问的是同一个公网域名（同一 Pages 部署）。")

doc.save(OUT)
print("已生成:", OUT)
